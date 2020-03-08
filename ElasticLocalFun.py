# -*- coding: utf-8 -*-
"""
Created on Sat May 18 15:11:44 2019

@author: nammabfc
"""

from elasticsearch import Elasticsearch
import os
import glob
from elasticsearch import helpers
import PyPDF2
import pandas as pd
import docx
from PyQt5 import QtWidgets,QtCore,uic
import time
import sys
import datetime
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *


from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import matplotlib.pyplot as plt

from PyQt5.QtCore import *

import ElasticLocal
import ElasticWaitFunc
import ElasticErrorFunc



class AppForm(QMainWindow):
    file=[]
    hit=[]
    def __init__(self,file=[],hit=[]):
        QMainWindow.__init__(self)
        self.setWindowTitle('Elasticsearch tool graph')
        self.file=file
        self.hit=hit
        self.create_menu()
        self.create_main_frame()
        self.create_status_bar()

        self.on_draw()

    def save_plot(self):
        file_choices = "PNG (*.png)|*.png"
        
        path, ext = QFileDialog.getSaveFileName(self, 
                        'Save file', '', 
                        file_choices)
        path = path.encode('utf-8')
        if not path[-4:] == file_choices[-4:].encode('utf-8'):
            path += file_choices[-4:].encode('utf-8')
        print(path)
        if path:
            self.canvas.print_figure(path.decode(), dpi=self.dpi)
            self.statusBar().showMessage('Saved to %s' % path, 2000)
    
    def on_about(self):
        msg = """ A demo of using PyQt with matplotlib:
        
         * Use the matplotlib navigation bar
         * Add values to the text box and press Enter (or click "Draw")
         * Show or hide the grid
         * Drag the slider to modify the width of the bars
         * Save the plot to a file using the File menu
         * Click on a bar to receive an informative message
        """
        QMessageBox.about(self, "About the demo", msg.strip())
    
    def on_pick(self, event):
        # The event received here is of the type
        # matplotlib.backend_bases.PickEvent
        #
        # It carries lots of information, of which we're using
        # only a small amount here.
        # 
        box_points = event.artist.get_bbox().get_points()
        msg = "You've clicked on a bar with coords:\n %s" % box_points
        
        QMessageBox.information(self, "Click!", msg)
    
    def on_draw(self):
        """ Redraws the figure
        """        
        # clear the axes and redraw the plot anew
        #
        self.axes.clear()        
        self.axes.grid(self.grid_cb.isChecked())
        
        self.axes.bar(
            self.file, 
            self.hit, 
            width=self.slider.value() / 100.0, 
            align='center', 
            alpha=0.44,
            picker=5)
        self.axes.set_ylabel("No of time word Found",fontsize=15)
        self.axes.set_xlabel("Books",fontsize=10)
        self.axes.set_title('Elastic Search using python')
        self.axes.set_xticklabels(self.file,{ 'fontsize':5}, rotation=30)
        self.canvas.draw()
    
    def create_main_frame(self):
        self.main_frame = QWidget()
        
        # Create the mpl Figure and FigCanvas objects. 
        # 5x4 inches, 100 dots-per-inch
        #
        self.dpi = 100
        self.fig = Figure((5.0, 20.0), dpi=self.dpi)
        self.canvas = FigureCanvas(self.fig)
        self.canvas.setParent(self.main_frame)
        
        # Since we have only one plot, we can use add_axes 
        # instead of add_subplot, but then the subplot
        # configuration tool in the navigation toolbar wouldn't
        # work.
        #
        self.axes = self.fig.add_subplot(111)
        
        # Bind the 'pick' event for clicking on one of the bars
        #
        self.canvas.mpl_connect('pick_event', self.on_pick)
        
        # Create the navigation toolbar, tied to the canvas
        #
        self.mpl_toolbar = NavigationToolbar(self.canvas, self.main_frame)
        
        # Other GUI controls
        # 
        
        self.grid_cb = QCheckBox("Show &Grid")
        self.grid_cb.setChecked(False)
        self.grid_cb.stateChanged.connect(self.on_draw)
        
        slider_label = QLabel('Bar width (%):')
        self.slider = QSlider(Qt.Horizontal)
        self.slider.setRange(1, 100)
        self.slider.setValue(20)
        self.slider.setTracking(True)
        self.slider.setTickPosition(QSlider.TicksBothSides)
        self.slider.valueChanged.connect(self.on_draw)
        
        #
        # Layout with box sizers
        # 
        hbox = QHBoxLayout()
        
        for w in [  self.grid_cb,
                    slider_label, self.slider]:
            hbox.addWidget(w)
            hbox.setAlignment(w, Qt.AlignVCenter)
        
        vbox = QVBoxLayout()
        vbox.addWidget(self.canvas)
        vbox.addWidget(self.mpl_toolbar)
        vbox.addLayout(hbox)
        
        self.main_frame.setLayout(vbox)
        self.setCentralWidget(self.main_frame)
    
    def create_status_bar(self):
        self.status_text = QLabel("Search Results")
        self.statusBar().addWidget(self.status_text, 1)
        
    def create_menu(self):        
        self.file_menu = self.menuBar().addMenu("&File")
        
        load_file_action = self.create_action("&Save plot",
            shortcut="Ctrl+S", slot=self.save_plot, 
            tip="Save the plot")
        quit_action = self.create_action("&Quit", slot=self.close, 
            shortcut="Ctrl+Q", tip="Close the application")
        
        self.add_actions(self.file_menu, 
            (load_file_action, None, quit_action))
        
        self.help_menu = self.menuBar().addMenu("&Help")
        about_action = self.create_action("&About", 
            shortcut='F1', slot=self.on_about, 
            tip='About the demo')
        
        self.add_actions(self.help_menu, (about_action,))

    def add_actions(self, target, actions):
        for action in actions:
            if action is None:
                target.addSeparator()
            else:
                target.addAction(action)

    def create_action(  self, text, slot=None, shortcut=None, 
                        icon=None, tip=None, checkable=False):
        action = QAction(text, self)
        if icon is not None:
            action.setIcon(QIcon(":/%s.png" % icon))
        if shortcut is not None:
            action.setShortcut(shortcut)
        if tip is not None:
            action.setToolTip(tip)
            action.setStatusTip(tip)
        if slot is not None:
            action.triggered.connect(slot)
        if checkable:
            action.setCheckable(True)
        return action

class MyWindow(QtWidgets.QMainWindow):
    es = Elasticsearch("http://localhost:9200")
    this_loc=1
    file=[]
    hit=[]
    def __init__(self, parent=None):
        QtWidgets.QMainWindow.__init__(self, parent)
        self.ui = ElasticLocal.Ui_MainWindow()
        self.ui.setupUi(self)
        
        self.ui.pushButton.clicked.connect(self.getfiles)
        self.ui.pushButton_2.clicked.connect(self.path)
        self.ui.pushButton_4.clicked.connect(self.search)
        self.ui.pushButton_6.clicked.connect(self.get)
        self.ui.pushButton_7.clicked.connect(self.download)
        
    def get(self, parent=None):
        self.vi = AppForm(self.file,self.hit)
        self.vi.show()
        
    def extractPdfFiles(self,files): 
        actions=[]
        for file in files:
            pdfFileObj = open(file,"rb")
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
            n_pages = pdfReader.numPages
            
            for i in range(n_pages):
                pageObj = pdfReader.getPage(i)
                this_text = pageObj.extractText()
                actions = actions+[
                        {
                                "_index": "data_science",
                                "_type": "books",
                                "_id": self.this_loc,
                                "_source": {
                                        "name":file,
                                        "page no": i+1,
                                        "content":this_text}
                         }
                        ]
                self.this_loc =self.this_loc+1
        return actions
    
    def extractTxtFiles(self,files):
        actions=[]
        for file in files:
            f=open(file,"r")
            cou=1
            for i in f:
                actions = actions+[
                        {
                            "_index": "data_science",
                            "_type": "books",
                            "_id": self.this_loc,
                            "_source": {
                                    "name":file,
                                    "line no": cou,
                                    "content":i}
                            }]
                cou=cou+1
                self.this_loc =self.this_loc+1
        return actions
    
    def extractExcleFiles(self,files):
        actions=[]
        for file in files:
            excel_f=pd.ExcelFile(file)
            for sheet in excel_f.sheet_names:
                df1=pd.read_excel(excel_f,sheet_name=sheet)
                for i in range(df1.shape[0]):
                    actions = actions+[
                      {
                            "_index": "data_science",
                            "_type": "books",
                            "_id": self.this_loc,
                            "_source": {
                                "name":file,
                                "sheet name":sheet,
                                "line no": i,
                                "content":str(df1.iloc[i])}
                      }
                ]
                    self.this_loc =self.this_loc+1
        return actions
    
    def extractWordFiles(self,files):
        actions=[]
        i=1
        for file in files:
            doc=docx.Document(file)
            for para in doc.paragraphs:
                actions=actions+[
                    {
                            "_index":"data_science",
                            "_type":"books",
                            "_id":self.this_loc,
                            "_source":{
                                    "name":file,
                                    "paragraph no":i,
                                    "content":para.text
                                }
                        }]
                i=i+1
                self.this_loc =self.this_loc+1
        return actions
    
    def path(self):
        path=self.ui.lineEdit.text()
        if path=='':
            self.error=ElasticErrorFunc.MyWindow()          
            self.error.show()
            self.error.ui.label.setText("Please enter path")
        
        else:
            path=path+"/"
            if os.path.isdir(path):
                self.wait=ElasticWaitFunc.MyWindow()
                self.wait.show()
                 
                os.chdir(path)
                files = glob.glob("*.pdf")
                df=self.extractPdfFiles(files)
                files = glob.glob("*.txt")
                df+=self.extractTxtFiles(files)
                files = glob.glob("*.xlsx")
                df+=self.extractExcleFiles(files)
                files =glob.glob("*.docx")
                df+=self.extractWordFiles(files)
                if self.es.indices.exists("data_science"):
                    self.es.indices.delete("data_science")
                helpers.bulk(self.es, df)
                self.wait.close()
                           
            else:
                self.error=ElasticErrorFunc.MyWindow()                 
                self.error.show()
                self.error.ui.label.setText("Invalid Path...Please enter valid path")
                  
                
    def search(self):
        if(self.es.indices.exists("data_science")):
            search_txt=self.ui.lineEdit_2.text()
            search_results = helpers.scan(self.es, index="data_science", query={"query": {"match": {"content":search_txt}}})
            self.file=[]
            self.hit=[]
            ans='<html>'
            cou=0
            title=''
            old=0
            for search_result in search_results:
                if(title!=str(search_result["_source"]["name"])):
                    if(cou>0):
                        ans+='<br>'+'<br>'+'<br>'+'<br>'
                        self.hit+=[cou-old]
                        old=cou
                    ans+='&emsp; '+'&emsp; '+"<b>"+str(search_result["_source"]["name"])+"</b>"+'<br>'+'<br>'
                    title=str(search_result["_source"]["name"])
                    self.file+=[title]
                for i in search_result["_source"].keys():
                    if(i=='name'):
                        pass
                    else:
                        ans+=i+':'+ str(search_result["_source"][i])+' &emsp;'
                
                ans+='<br>'+'<br>'
                cou+=1
            self.hit+=[cou-old]   
            ans+='</html>'
            self.ui.lineEdit_3.setText(str(cou))
            self.ui.textEdit.setText(ans)
        else:
            self.error=ElasticErrorFunc.MyWindow()                 
            self.error.show()
            self.error.ui.label.setText("Please enter path and then enter search text")
        
    def getfiles(self):
        fileName = str(QtWidgets.QFileDialog.getExistingDirectory(self, "Select Directory"))
        self.ui.lineEdit.setText(fileName)

    def download(self):
        content=self.ui.textEdit.toPlainText()
        if(content==''):
            self.error=ElasticErrorFunc.MyWindow()                 
            self.error.show()
            self.error.ui.label.setText("Output box is empty...Enter valid search text to download")
        else:
            x=datetime.datetime.now()
            document=docx.Document()
            document.add_heading('Report',0)    
            p=document.add_paragraph(content)
            document.save('report '+str(x.strftime("%d"))+'-'+str(x.strftime("%m"))+'-'+str(x.strftime("%Y"))
            +' '+str(x.strftime("%H"))+' '+str(x.strftime("%M"))+' '+str(x.strftime("%S"))+'.docx')


if __name__ == '__main__':
    
    app = QtWidgets.QApplication.instance()
    if app is None:
        app = QtWidgets.QApplication(sys.argv)
    app.lastWindowClosed.connect(app.quit)
    window = MyWindow()
    window.show()
    sys.exit(app.exec_())