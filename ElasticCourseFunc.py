# -*- coding: utf-8 -*-
"""
Created on Sun May 26 10:05:57 2019

@author: Vector30
"""
import datetime,docx
from docx import Document
from PyQt5 import QtWidgets,QtCore
import ElasticCourse
class MyWindow(QtWidgets.QMainWindow):
     def __init__(self, parent=None):
        QtWidgets.QMainWindow.__init__(self)
        self.ui = ElasticCourse.Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(self.download)
          
        
     def download(self):
        x=datetime.datetime.now()
        document=docx.Document()
        document.add_heading(self.ui.label.text(),0)
        a=document.add_paragraph()
        p=a.add_run("Institute : ")
        font=p.font
        font.bold=True
        a.add_run(self.ui.label_2.text())
        b=document.add_paragraph()
        p=b.add_run("Professor : ")
        font=p.font
        font.bold=True
        b.add_run(self.ui.label_3.text())
        p=document.add_paragraph().add_run("Description : ")
        font=p.font
        font.bold=True            
        p=document.add_paragraph(self.ui.textEdit_2.toPlainText())
        p=document.add_paragraph().add_run("Summary : ")
        font=p.font
        font.bold=True
        table=document.add_table(rows=8 , cols=3)
        table.autofit=True
        cells=table.rows[0].cells
        cells[0].text="1."
        cells[1].text="Course Status : "
        cells[2].text=self.ui.label_13.text()
        
        cells=table.rows[1].cells
        cells[0].text="2."
        cells[1].text="Course Type : "
        cells[2].text=self.ui.label_14.text()
        
        cells=table.rows[2].cells
        cells[0].text="3."
        cells[1].text="Duration : "
        cells[2].text=self.ui.label_15.text()
        
        cells=table.rows[3].cells
        cells[0].text="4."
        cells[1].text="Start Date : "
        cells[2].text=self.ui.label_16.text()
        
        cells=table.rows[4].cells
        cells[0].text="5."
        cells[1].text="End Date : "
        cells[2].text=self.ui.label_17.text()
        
        cells=table.rows[5].cells
        cells[0].text="6."
        cells[1].text="Exam Date : "
        cells[2].text=self.ui.label_18.text()
        
        cells=table.rows[6].cells
        cells[0].text="7."
        cells[1].text="Category : "
        cells[2].text=self.ui.label_19.text()
        
        cells=table.rows[7].cells
        cells[0].text="8."
        cells[1].text="Level : "
        cells[2].text=self.ui.label_20.text()
        
        p=document.add_paragraph().add_run("Course Layout : ")
        
        p.font.bold=True
        print(font.bold)
        p=document.add_paragraph(self.ui.textEdit.toPlainText())
        document.save(self.ui.label.text()+" "+str(x.strftime("%d"))+'-'+str(x.strftime("%m"))+'-'+str(x.strftime("%Y"))
        +' '+str(x.strftime("%H"))+' '+str(x.strftime("%M"))+' '+str(x.strftime("%S"))+'.docx')
        
        
